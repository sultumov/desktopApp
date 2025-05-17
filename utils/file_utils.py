"""Утилиты для работы с файлами."""

import os
import json
import logging
from pathlib import Path
from typing import Optional, Dict, Any
from PyQt6.QtWidgets import QFileDialog, QMessageBox
from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# Настройка логгера
logger = logging.getLogger(__name__)

def save_text_to_file(parent, text, title="Сохранить файл", default_path="", file_filter="Текстовые файлы (*.txt);;Все файлы (*.*)"):
    """Сохраняет текст в файл с диалогом выбора имени файла.

    Args:
        parent: Родительский виджет для диалога
        text: Текст для сохранения
        title: Заголовок диалога сохранения
        default_path: Путь по умолчанию
        file_filter: Фильтр типов файлов

    Returns:
        Кортеж (успех: bool, сообщение: str)
    """
    if not text:
        return False, "Нет текста для сохранения"

    file_name, _ = QFileDialog.getSaveFileName(
        parent,
        title,
        default_path,
        file_filter
    )

    if not file_name:
        return False, "Отменено пользователем"

    try:
        with open(file_name, 'w', encoding='utf-8') as f:
            f.write(text)
        return True, f"Данные сохранены в файл {file_name}"
    except Exception as e:
        logger.error(f"Ошибка при сохранении файла {file_name}: {str(e)}")
        return False, f"Ошибка при сохранении файла: {str(e)}"

def ensure_dir_exists(directory):
    """Создаёт директорию, если она не существует.

    Args:
        directory: Путь к директории

    Returns:
        True, если директория существует или была создана успешно
    """
    if not os.path.exists(directory):
        try:
            os.makedirs(directory)
            logger.info(f"Создана директория {directory}")
            return True
        except Exception as e:
            logger.error(f"Ошибка при создании директории {directory}: {str(e)}")
            return False
    return True

def export_article_to_file(parent, article, title="Экспортировать статью", default_path="", 
                          file_filter="Текстовые файлы (*.txt);;PDF файлы (*.pdf);;Word документы (*.docx);;Все файлы (*.*)"):
    """Экспортирует информацию о статье в файл.

    Args:
        parent: Родительский виджет для диалога
        article: Объект статьи для экспорта
        title: Заголовок диалога экспорта
        default_path: Путь по умолчанию
        file_filter: Фильтр типов файлов

    Returns:
        Кортеж (успех: bool, сообщение: str)
    """
    if not article:
        return False, "Статья не выбрана"

    file_name, selected_filter = QFileDialog.getSaveFileName(
        parent,
        title,
        default_path,
        file_filter
    )

    if not file_name:
        return False, "Отменено пользователем"

    try:
        if file_name.lower().endswith('.pdf'):
            return export_to_pdf(file_name, article)
        elif file_name.lower().endswith('.docx'):
            return export_to_docx(file_name, article)
        else:
            return export_to_txt(file_name, article)
    except Exception as e:
        logger.error(f"Ошибка при экспорте статьи в файл {file_name}: {str(e)}")
        return False, f"Ошибка при экспорте статьи: {str(e)}"

def export_to_txt(file_name, article):
    """Экспортирует статью в текстовый файл."""
    content = format_article_content(article)
    with open(file_name, 'w', encoding='utf-8') as f:
        f.write(content)
    return True, f"Статья экспортирована в {file_name}"

def export_to_pdf(file_name, article):
    """Экспортирует статью в PDF файл."""
    doc = SimpleDocTemplate(file_name, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Создаем стили
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12
    )
    
    normal_style = styles['Normal']
    
    # Создаем элементы документа
    elements = []
    
    # Заголовок
    elements.append(Paragraph(article.title, title_style))
    elements.append(Spacer(1, 12))
    
    # Авторы
    elements.append(Paragraph("Авторы:", heading_style))
    elements.append(Paragraph(', '.join(article.authors), normal_style))
    elements.append(Spacer(1, 12))
    
    # Дата публикации
    elements.append(Paragraph("Дата публикации:", heading_style))
    elements.append(Paragraph(article.published.strftime('%d.%m.%Y'), normal_style))
    elements.append(Spacer(1, 12))
    
    # Категории
    elements.append(Paragraph("Категории:", heading_style))
    elements.append(Paragraph(', '.join(article.categories), normal_style))
    elements.append(Spacer(1, 12))
    
    # DOI и URL
    if article.doi:
        elements.append(Paragraph("DOI:", heading_style))
        elements.append(Paragraph(article.doi, normal_style))
        elements.append(Spacer(1, 12))
    
    elements.append(Paragraph("URL:", heading_style))
    elements.append(Paragraph(article.url, normal_style))
    elements.append(Spacer(1, 12))
    
    # Аннотация
    elements.append(Paragraph("Аннотация:", heading_style))
    elements.append(Paragraph(article.summary, normal_style))
    
    # Создаем PDF
    doc.build(elements)
    return True, f"Статья экспортирована в PDF: {file_name}"

def export_to_docx(file_name, article):
    """Экспортирует статью в DOCX файл."""
    doc = Document()
    
    # Заголовок
    doc.add_heading(article.title, level=1)
    
    # Авторы
    doc.add_heading("Авторы:", level=2)
    doc.add_paragraph(', '.join(article.authors))
    
    # Дата публикации
    doc.add_heading("Дата публикации:", level=2)
    doc.add_paragraph(article.published.strftime('%d.%m.%Y'))
    
    # Категории
    doc.add_heading("Категории:", level=2)
    doc.add_paragraph(', '.join(article.categories))
    
    # DOI и URL
    if article.doi:
        doc.add_heading("DOI:", level=2)
        doc.add_paragraph(article.doi)
    
    doc.add_heading("URL:", level=2)
    doc.add_paragraph(article.url)
    
    # Аннотация
    doc.add_heading("Аннотация:", level=2)
    doc.add_paragraph(article.summary)
    
    # Сохраняем документ
    doc.save(file_name)
    return True, f"Статья экспортирована в DOCX: {file_name}"

def format_article_content(article):
    """Форматирует содержимое статьи для текстового файла."""
    content = f"Название: {article.title}\n"
    content += f"Авторы: {', '.join(article.authors)}\n"
    content += f"Дата публикации: {article.published.strftime('%d.%m.%Y')}\n"
    content += f"Категории: {', '.join(article.categories)}\n"
    content += f"DOI: {article.doi or 'Не указан'}\n"
    content += f"URL: {article.url}\n\n"
    content += "Аннотация:\n"
    content += f"{article.summary}\n"
    return content

def open_file(file_path):
    """Открывает файл в ассоциированной программе.

    Args:
        file_path: Путь к файлу

    Returns:
        Кортеж (успех: bool, сообщение: str)
    """
    if not os.path.exists(file_path):
        return False, f"Файл {file_path} не найден"

    try:
        os.startfile(file_path)
        return True, f"Файл {file_path} открыт"
    except Exception as e:
        logger.error(f"Ошибка при открытии файла {file_path}: {str(e)}")
        return False, f"Ошибка при открытии файла: {str(e)}"

def confirm_file_action(parent, title, message, default_yes=False):
    """Запрашивает подтверждение для действия с файлом.

    Args:
        parent: Родительский виджет для диалога
        title: Заголовок диалога
        message: Текст сообщения
        default_yes: True, если кнопка "Да" должна быть по умолчанию

    Returns:
        True, если пользователь подтвердил действие
    """
    default_button = QMessageBox.StandardButton.Yes if default_yes else QMessageBox.StandardButton.No
    
    reply = QMessageBox.question(
        parent,
        title,
        message,
        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        default_button
    )
    
    return reply == QMessageBox.StandardButton.Yes 